import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt

# Load the data from the CSV file
data = pd.read_csv("GNG4120.csv")

# Drop the 'Student ID' column
data.drop(columns=['Student ID'], inplace=True)

# Calculate observed values (standard deviation)
observed_values = data.std(axis=0)

# Calculate predicted values (mean)
predicted_values = data.mean(axis=0)

# Create a Word document
doc = Document()
doc.add_heading('Statistical Analysis', level=1)

# Add mean and standard deviation table for all questions
doc.add_heading('Mean and Standard Deviation', level=2)
mean_std_table = doc.add_table(rows=len(observed_values)+1, cols=2)
hdr_cells = mean_std_table.rows[0].cells
hdr_cells[0].text = 'Question'
hdr_cells[1].text = 'Mean ± Std Deviation'

for i, (question, observed_val) in enumerate(zip(observed_values.index, observed_values)):
    row_cells = mean_std_table.rows[i+1].cells
    row_cells[0].text = question
    row_cells[1].text = f"{predicted_values[i]:.2f} ± {observed_val:.2f}"

# Ensure all columns are considered as questions
num_questions = len(data.columns)

# Plot all questions with both observed and predicted counts
fig, ax = plt.subplots(figsize=(12, 6))
labels = data.columns
x = range(1, len(labels)+1)  # Changing x-axis to start from 1
bar_width = 0.35
bars1 = ax.bar(x, observed_values, width=bar_width, label='Standard Deviation')
bars2 = ax.bar([i + bar_width for i in x], predicted_values, width=bar_width, label='Mean', alpha=0.5)
ax.set_xticks(x)
ax.set_xticklabels(x, fontsize=16)  # Setting x-axis labels to question numbers
ax.legend(fontsize=16)
ax.set_xlabel('Question Numbers', fontsize=16)  # Changing x-axis label to "Question Numbers"
ax.set_ylabel('Answers (Strongly Disagree-> Strongly Agree)', fontsize=16)
ax.tick_params(axis='both', which='major', labelsize=16)
ax.set_ylim(0, 5)  # Set y-axis limit to 5

# Annotate each bar with the corresponding value
for bar in bars1 + bars2:
    height = bar.get_height()
    ax.annotate(f'{height:.2f}',
                xy=(bar.get_x() + bar.get_width() / 2, height),
                xytext=(0, 3),  # 3 points vertical offset
                textcoords="offset points",
                ha='center', va='bottom', fontsize=16)

plt.tight_layout()

# Add the chart to the Word document
doc.add_picture('observed_vs_predicted_values_all_questions.png', width=Inches(6))

# Save the document
doc.save('statistical_analysis.docx')

# Show the figure
plt.show()
